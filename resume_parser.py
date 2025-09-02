#!/usr/bin/env python3
"""
Resume Parser Module
Extracts raw text from various resume formats (PDF, Word, etc.)
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse resumes in various formats and extract raw text."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a resume file and extract raw text with metadata.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing parsed text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._parse_word(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_path.suffix}")
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF resume using multiple methods for better extraction."""
        logger.info(f"Parsing PDF: {file_path}")
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                text_content = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text.strip())
                
                if any(text_content):
                    return {
                        'file_path': str(file_path),
                        'file_type': 'pdf',
                        'text': '\n\n'.join(text_content),
                        'pages': len(pdf.pages),
                        'parser_method': 'pdfplumber'
                    }
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text.strip())
                
                return {
                    'file_path': str(file_path),
                    'file_type': 'pdf',
                    'text': '\n\n'.join(text_content),
                    'pages': len(pdf_reader.pages),
                    'parser_method': 'PyPDF2'
                }
        except Exception as e:
            logger.error(f"PyPDF2 parsing failed: {e}")
            raise
    
    def _parse_word(self, file_path: Path) -> Dict[str, Any]:
        """Parse Word document resume."""
        logger.info(f"Parsing Word document: {file_path}")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return {
                'file_path': str(file_path),
                'file_type': 'word',
                'text': '\n\n'.join(text_content),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'parser_method': 'python-docx'
            }
        except Exception as e:
            logger.error(f"Word document parsing failed: {e}")
            raise
    
    def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text resume."""
        logger.info(f"Parsing text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            return {
                'file_path': str(file_path),
                'file_type': 'text',
                'text': text_content,
                'parser_method': 'text'
            }
        except Exception as e:
            logger.error(f"Text file parsing failed: {e}")
            raise
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with LLM processing
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        return text.strip()
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Attempt to identify and extract common resume sections.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            Dictionary with section names and content
        """
        sections = {}
        
        # Common section headers
        section_patterns = {
            'contact': r'(?i)(contact|personal|address|phone|email)',
            'summary': r'(?i)(summary|profile|objective|about)',
            'experience': r'(?i)(experience|work\s+history|employment)',
            'education': r'(?i)(education|academic|degree)',
            'skills': r'(?i)(skills|technical\s+skills|competencies)',
            'projects': r'(?i)(projects|portfolio|achievements)',
            'certifications': r'(?i)(certifications|certificates|licenses)',
            'languages': r'(?i)(languages|language\s+skills)',
            'interests': r'(?i)(interests|hobbies|activities)'
        }
        
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_lower):
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    section_found = True
                    break
            
            if not section_found:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections


def main():
    """Test the resume parser."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Parse resume files')
    parser.add_argument('file_path', help='Path to resume file')
    parser.add_argument('--clean', action='store_true', help='Clean extracted text')
    parser.add_argument('--sections', action='store_true', help='Extract sections')
    
    args = parser.parse_args()
    
    try:
        resume_parser = ResumeParser()
        result = resume_parser.parse_resume(args.file_path)
        
        print(f"Successfully parsed: {result['file_path']}")
        print(f"File type: {result['file_type']}")
        print(f"Parser method: {result['parser_method']}")
        
        if args.clean:
            cleaned_text = resume_parser.clean_text(result['text'])
            print(f"\nCleaned text (first 500 chars):\n{cleaned_text[:500]}...")
        
        if args.sections:
            cleaned_text = resume_parser.clean_text(result['text'])
            sections = resume_parser.extract_sections(cleaned_text)
            print(f"\nExtracted sections:")
            for section, content in sections.items():
                print(f"\n{section.upper()}:")
                print(f"{content[:200]}...")
        
    except Exception as e:
        print(f"Error: {e}")
        return 1
    
    return 0


if __name__ == '__main__':
    exit(main())

